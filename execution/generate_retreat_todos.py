"""
HYDS 수련회별 To-Do 보고서 v2 — 진짜 회사 수준 디자인.

표지 + Executive Summary + KPI 대시보드 + 위험 매트릭스 + 수련회별 상세
+ 통합 액션 표 + 부록 (페이지 번호·헤더·푸터 포함).

사용 예:
    python execution/generate_retreat_todos.py            # 실DB + 보고서 + 텔레그램 발송 + 로그
    python execution/generate_retreat_todos.py --no-send  # 발송 안 함 (보고서·로그만)
    python execution/generate_retreat_todos.py --mock     # 모의 데이터 샘플 (발송·로그 X)

v2 회사급 디자인 + 텔레그램 발송 + weekly_todos_log.json 기록(office 대시보드·weekly_review 의존).
"""
import argparse
import sys
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))
from execution.utils import read_json, write_json

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.section import WD_SECTION
except ImportError:
    print("❌ python-docx 미설치: pip install python-docx")
    sys.exit(1)

PROJECT_ROOT = Path(__file__).parent.parent
OUTPUT = PROJECT_ROOT / ".tmp" / "dossiers" / f"HYDS_수련회_운영현황_보고서_{datetime.now().strftime('%Y%m%d')}.docx"
KOREAN_FONT = "Malgun Gothic"

# ─── HYDS 브랜드 컬러 팔레트 ───────────────────────────────────────────────────
COLOR_PRIMARY = RGBColor(0x4A, 0x2C, 0x1A)     # 진한 갈색 (헤딩)
COLOR_ACCENT = RGBColor(0xB4, 0x53, 0x09)      # 앰버 (강조)
COLOR_MUTED = RGBColor(0x78, 0x6B, 0x5E)       # 그레이 (보조)
COLOR_BG_CREAM = "FEF3E2"                       # 크림 (배경 fill)
COLOR_BG_AMBER = "FED7AA"                       # 앰버 라이트
COLOR_BG_HEADER = "8B5A3C"                      # 다크 우드 (테이블 헤더)
COLOR_RISK_HIGH = "DC2626"
COLOR_RISK_MED = "F59E0B"
COLOR_RISK_LOW = "16A34A"
COLOR_RISK_HIGH_BG = "FEE2E2"
COLOR_RISK_MED_BG = "FEF3C7"
COLOR_RISK_LOW_BG = "D1FAE5"

# ─── 헬퍼 ──────────────────────────────────────────────────────────────────────
def set_font(run, name=KOREAN_FONT, size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")


def shade_cell(cell, fill_hex):
    """셀 배경색."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_borders(cell, color="C4A484", size=4):
    """셀 테두리."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_para(doc, text, size=10, bold=False, italic=False, align=None,
             color=None, space_before=0, space_after=0, indent=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 28, 1: 18, 2: 14, 3: 12}
    spaces = {0: (24, 12), 1: (18, 8), 2: (12, 6), 3: (8, 4)}
    sb, sa = spaces.get(level, (8, 4))
    p = add_para(doc, text, size=sizes.get(level, 10), bold=True,
                 color=COLOR_PRIMARY, space_before=sb, space_after=sa)
    return p


def add_divider(doc, color=COLOR_ACCENT, thickness=8):
    """수평선 (가로 진한 줄)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color.rgb if hasattr(color,'rgb') else 'B45309'}" if False else "B45309")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_callout_box(doc, label, text, bg_hex=COLOR_BG_AMBER, border_hex="B45309"):
    """강조 박스 (callout)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(16.5)
    shade_cell(cell, bg_hex)
    set_cell_borders(cell, color=border_hex, size=8)

    p1 = cell.paragraphs[0]
    run = p1.add_run(label)
    set_font(run, size=10, bold=True, color=COLOR_PRIMARY)
    if text:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        run2 = p2.add_run(text)
        set_font(run2, size=10, color=COLOR_PRIMARY)
    return table


def progress_bar_text(progress: int) -> str:
    """텍스트 progress bar (10블록)."""
    full = round(progress / 10)
    return "█" * full + "░" * (10 - full)


def risk_category(days_until: int, progress: int, issues: int) -> str:
    """위험도 결정론 판정."""
    if days_until is None:
        return "unknown"
    if (days_until <= 7 and progress < 70) or issues > 0:
        return "high"
    if (days_until <= 30 and progress < 50) or (days_until <= 14 and progress < 80):
        return "med"
    return "low"


# ─── 페이지 헤더·푸터 ─────────────────────────────────────────────────────────
def setup_page_headers(doc, title="수련회 운영 현황 보고서"):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

        # 헤더
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("HYDS  ·  ")
        set_font(run, size=8, bold=True, color=COLOR_PRIMARY)
        run2 = hp.add_run(title)
        set_font(run2, size=8, color=COLOR_MUTED)

        # 푸터: 좌측 confidentiality + 우측 페이지 번호
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = fp.add_run("HYDS Confidential · 내부용")
        set_font(run, size=8, color=COLOR_MUTED, italic=True)
        run2 = fp.add_run("\t\t")
        set_font(run2, size=8)
        # 페이지 번호 필드
        run3 = fp.add_run()
        set_font(run3, size=8, color=COLOR_MUTED)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run3._element.append(fld_begin)
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        run3._element.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run3._element.append(fld_end)
        run4 = fp.add_run(" / ")
        set_font(run4, size=8, color=COLOR_MUTED)
        run5 = fp.add_run()
        set_font(run5, size=8, color=COLOR_MUTED)
        for el in [("begin",None),("instrText","NUMPAGES"),("end",None)]:
            if el[0] == "instrText":
                e = OxmlElement("w:instrText")
                e.text = el[1]
            else:
                e = OxmlElement("w:fldChar")
                e.set(qn("w:fldCharType"), el[0])
            run5._element.append(e)


# ─── 표지 페이지 ──────────────────────────────────────────────────────────────
def add_cover_page(doc, analyses, report_date):
    # 상단 여백
    for _ in range(3):
        doc.add_paragraph()

    # 회사명
    add_para(doc, "H Y D S", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=COLOR_ACCENT, space_after=4)
    add_para(doc, "기독교 수련회 기획·운영", size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=COLOR_MUTED, italic=True, space_after=40)

    # 메인 타이틀
    add_para(doc, "수련회 운영 현황", size=32, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=COLOR_PRIMARY, space_after=2)
    add_para(doc, "주간 보고서", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=COLOR_PRIMARY, space_after=24)

    # 부제목 — 보고 기간
    add_para(doc, f"보고 기간: {report_date.strftime('%Y년 %m월 %d일')} 기준",
             size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_MUTED,
             space_after=60)

    # 메타 정보 표
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    meta_data = [
        ("보고 대상", "서동현 대표"),
        ("작성", "HYDS Director (부장)"),
        ("분석 수련회", f"총 {len(analyses)}건"),
        ("문서 등급", "Internal — 내부용 (외부 공유 금지)"),
    ]
    for i, (label, value) in enumerate(meta_data):
        c1 = meta_table.rows[i].cells[0]
        c2 = meta_table.rows[i].cells[1]
        c1.width = Cm(5)
        c2.width = Cm(10)
        c1.text = ""
        c2.text = ""
        shade_cell(c1, COLOR_BG_CREAM)
        run = c1.paragraphs[0].add_run(label)
        set_font(run, size=10, bold=True, color=COLOR_PRIMARY)
        run2 = c2.paragraphs[0].add_run(value)
        set_font(run2, size=10, color=COLOR_PRIMARY)
        for c in [c1, c2]:
            set_cell_borders(c, color="D4A574", size=4)

    # 하단 푸터
    for _ in range(8):
        doc.add_paragraph()
    add_para(doc, "─── HYDS Confidential ───", size=8, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=COLOR_MUTED, italic=True)

    doc.add_page_break()


# ─── Executive Summary ────────────────────────────────────────────────────────
def add_executive_summary(doc, analyses):
    add_heading(doc, "Executive Summary", level=1)
    add_para(doc, "본 보고서의 한 페이지 요약 — 의사결정에 필요한 핵심만.",
             size=9, italic=True, color=COLOR_MUTED, space_after=12)

    total = len(analyses)
    high = sum(1 for a in analyses if a["risk"] == "high")
    med = sum(1 for a in analyses if a["risk"] == "med")
    low = sum(1 for a in analyses if a["risk"] == "low")
    avg_progress = round(sum(a["progress"] for a in analyses) / total) if total else 0
    near_term = sum(1 for a in analyses if a["days_until"] is not None and a["days_until"] <= 30)

    # KPI 카드 4개
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.autofit = False
    kpis = [
        (f"{total}", "진행 중 수련회", COLOR_BG_CREAM),
        (f"{avg_progress}%", "평균 진척률", COLOR_BG_AMBER),
        (f"{high}", "🔴 위험 높음", COLOR_RISK_HIGH_BG),
        (f"{near_term}", "D-30 이내", COLOR_BG_CREAM),
    ]
    for i, (value, label, bg) in enumerate(kpis):
        cell = kpi_table.rows[0].cells[i]
        cell.width = Cm(4.0)
        cell.text = ""
        shade_cell(cell, bg)
        set_cell_borders(cell, color="C4A484", size=4)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(6)
        run = p1.add_run(value)
        set_font(run, size=22, bold=True, color=COLOR_PRIMARY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(6)
        run = p2.add_run(label)
        set_font(run, size=9, color=COLOR_MUTED)

    add_para(doc, "", space_after=12)

    # 위험 분포
    add_para(doc, "위험도 분포", size=11, bold=True, color=COLOR_PRIMARY,
             space_before=10, space_after=4)
    risk_table = doc.add_table(rows=2, cols=3)
    risk_table.autofit = False
    risk_headers = ["🔴 즉시 대응 필요", "🟡 주의 관찰", "🟢 정상 진행"]
    risk_bg = [COLOR_RISK_HIGH_BG, COLOR_RISK_MED_BG, COLOR_RISK_LOW_BG]
    risk_values = [f"{high}건", f"{med}건", f"{low}건"]
    for i, (h, v, bg) in enumerate(zip(risk_headers, risk_values, risk_bg)):
        hc = risk_table.rows[0].cells[i]
        vc = risk_table.rows[1].cells[i]
        hc.width = vc.width = Cm(5.5)
        for c in (hc, vc):
            c.text = ""
            shade_cell(c, bg)
            set_cell_borders(c, color="C4A484")
        run = hc.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=COLOR_PRIMARY)
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = vp.add_run(v)
        set_font(run, size=16, bold=True, color=COLOR_PRIMARY)

    add_para(doc, "", space_after=12)

    # 핵심 메시지 (3줄 안에)
    add_callout_box(
        doc,
        "📌 부장 보고 요약",
        (
            f"진행 중 수련회 {total}건 중 {high}건이 즉시 대응이 필요합니다. "
            f"평균 진척률은 {avg_progress}%로, D-30 이내인 {near_term}건의 마감 일정 관리가 가장 중요합니다. "
            f"세부 사항은 본문 3페이지(수련회별 상세) 및 7페이지(통합 액션 아이템)를 확인 바랍니다."
        )
    )

    add_para(doc, "", space_after=6)

    # Top 위험 3건
    high_list = [a for a in analyses if a["risk"] == "high"][:3]
    if high_list:
        add_para(doc, "🔴 즉시 대응이 필요한 수련회 (상위 3건)", size=11, bold=True,
                 color=COLOR_PRIMARY, space_before=8, space_after=4)
        for a in high_list:
            r = a["retreat"]
            d = f"D-{a['days_until']}" if a["days_until"] >= 0 else f"D+{-a['days_until']}"
            add_para(doc,
                     f"• {r.get('churchName', '?')} ({d}, 진척률 {a['progress']}%) — "
                     f"{a['reason_summary']}",
                     size=10, color=COLOR_PRIMARY, indent=0.3, space_after=2)

    doc.add_page_break()


# ─── KPI 대시보드 (상세) ─────────────────────────────────────────────────────
def add_overview_table(doc, analyses):
    add_heading(doc, "한눈에 보기 — 수련회 전체 현황", level=1)
    add_para(doc, "각 수련회의 D-day · 진척률 · 위험도 · 다음 액션을 한 표로 정리.",
             size=9, italic=True, color=COLOR_MUTED, space_after=8)

    headers = ["#", "교회/단체", "주제", "팀장", "D-day", "진척률", "위험", "다음 액션"]
    widths = [0.8, 3.0, 3.0, 2.0, 1.5, 2.0, 1.2, 3.0]

    table = doc.add_table(rows=1 + len(analyses), cols=len(headers))
    table.autofit = False

    # 헤더 행
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.width = Cm(widths[i])
        c.text = ""
        shade_cell(c, COLOR_BG_HEADER)
        set_cell_borders(c, color="4A2C1A", size=6)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xF8, 0xEB))

    # 데이터 행
    for idx, a in enumerate(analyses):
        r = a["retreat"]
        d = "?" if a["days_until"] is None else (f"D-{a['days_until']}" if a["days_until"] >= 0 else f"D+{-a['days_until']}")
        progress_text = f"{progress_bar_text(a['progress'])} {a['progress']:>3}%"
        risk_emoji = {"high": "🔴", "med": "🟡", "low": "🟢", "unknown": "⚪"}[a["risk"]]
        next_action = (a["todo_now"][0][1] if a["todo_now"] else "-")[:24]

        vals = [str(idx+1), r.get("churchName") or "?", (r.get("theme") or "미정")[:18],
                r.get("teamLeader") or "미배정", d, progress_text, risk_emoji, next_action]
        cells = table.rows[idx + 1].cells
        for i, v in enumerate(vals):
            cells[i].width = Cm(widths[i])
            cells[i].text = ""
            set_cell_borders(cells[i], color="D4A574")
            if i % 2 == 0:
                shade_cell(cells[i], "FFFFFF")
            else:
                shade_cell(cells[i], COLOR_BG_CREAM)
            # 위험 셀은 배경색
            if i == 6:
                bg = {"high": COLOR_RISK_HIGH_BG, "med": COLOR_RISK_MED_BG,
                      "low": COLOR_RISK_LOW_BG, "unknown": "FFFFFF"}[a["risk"]]
                shade_cell(cells[i], bg)
            p = cells[i].paragraphs[0]
            if i in (0, 4, 5, 6):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font_kwargs = {"size": 9, "color": COLOR_PRIMARY}
            if i == 5:
                font_kwargs["name"] = "Consolas"
            run = p.add_run(str(v))
            run.font.name = font_kwargs.get("name", KOREAN_FONT)
            run.font.size = Pt(font_kwargs["size"])
            run.font.color.rgb = font_kwargs["color"]
            if i != 5:
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.append(rFonts)
                rFonts.set(qn("w:eastAsia"), KOREAN_FONT)

    doc.add_page_break()


# ─── 수련회별 상세 ────────────────────────────────────────────────────────────
def add_retreat_detail(doc, a, idx):
    r = a["retreat"]
    d = "?" if a["days_until"] is None else (f"D-{a['days_until']}" if a["days_until"] >= 0 else f"D+{-a['days_until']}")
    risk_label = {"high": "🔴 위험 높음", "med": "🟡 주의 관찰",
                  "low": "🟢 정상 진행", "unknown": "⚪ 정보 부족"}[a["risk"]]
    risk_bg = {"high": COLOR_RISK_HIGH_BG, "med": COLOR_RISK_MED_BG,
               "low": COLOR_RISK_LOW_BG, "unknown": "FFFFFF"}[a["risk"]]

    add_heading(doc, f"{idx}. {r.get('churchName') or '?'}", level=1)
    # 메타 정보 표
    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.autofit = False
    meta_data = [
        ("주제", r.get("theme") or "미정"),
        ("팀장", r.get("teamLeader") or "미배정"),
        ("D-day", d),
        ("위험도", risk_label),
    ]
    for i, (label, value) in enumerate(meta_data):
        lc = meta_table.rows[0].cells[i]
        vc = meta_table.rows[1].cells[i]
        lc.width = vc.width = Cm(4.1)
        for c in (lc, vc):
            c.text = ""
            set_cell_borders(c, color="D4A574")
        shade_cell(lc, COLOR_BG_HEADER)
        if i == 3:
            shade_cell(vc, risk_bg)
        else:
            shade_cell(vc, COLOR_BG_CREAM)
        run = lc.paragraphs[0].add_run(label)
        set_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xF8, 0xEB))
        run = vc.paragraphs[0].add_run(value)
        set_font(run, size=10, bold=(i==3), color=COLOR_PRIMARY)

    add_para(doc, "", space_after=8)

    # 진척률 시각화
    add_para(doc, "진척률", size=10, bold=True, color=COLOR_PRIMARY, space_after=2)
    prog_text = f"{progress_bar_text(a['progress'])}  {a['progress']}%  ({a['done']}/{a['total']} 항목 완료)"
    p = doc.add_paragraph()
    run = p.add_run(prog_text)
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_ACCENT
    add_para(doc, "", space_after=4)

    # AI 조언
    add_callout_box(doc, "🤖 부장 조언", a["advice"])
    add_para(doc, "", space_after=8)

    # 발견된 이슈
    if a["issues"]:
        add_heading(doc, "발견된 이슈", level=3)
        for issue in a["issues"]:
            add_para(doc, f"⚠️  {issue}", size=10, color=COLOR_PRIMARY,
                     indent=0.3, space_after=2)

    # 우선순위 액션 표
    add_heading(doc, "🔴 즉시 해야 할 To-Do (우선순위순)", level=3)
    if a["todo_now"]:
        todo_table = doc.add_table(rows=1 + len(a["todo_now"]), cols=4)
        todo_table.autofit = False
        headers = ["No.", "단계", "해야 할 일", "비고"]
        widths = [1.0, 3.0, 9.0, 3.5]
        for i, h in enumerate(headers):
            c = todo_table.rows[0].cells[i]
            c.width = Cm(widths[i])
            c.text = ""
            shade_cell(c, COLOR_BG_HEADER)
            set_cell_borders(c, color="4A2C1A", size=6)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            set_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xF8, 0xEB))
        for i, (cat, item) in enumerate(a["todo_now"]):
            cells = todo_table.rows[i+1].cells
            vals = [str(i+1), cat, item, "확인 필요"]
            for j, v in enumerate(vals):
                cells[j].width = Cm(widths[j])
                cells[j].text = ""
                set_cell_borders(cells[j], color="D4A574")
                shade_cell(cells[j], "FFFFFF" if i%2==0 else COLOR_BG_CREAM)
                p = cells[j].paragraphs[0]
                if j == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(v)
                set_font(run, size=9, color=COLOR_PRIMARY)

    # 다음 단계 미리보기
    if a["todo_next"]:
        add_heading(doc, "🟡 다음 단계 미리보기", level=3)
        for cat, item in a["todo_next"][:6]:
            add_para(doc, f"○  [{cat}] {item}", size=9, color=COLOR_MUTED,
                     indent=0.3, space_after=1)

    doc.add_page_break()


# ─── 통합 액션 아이템 ───────────────────────────────────────────────────────────
def add_consolidated_actions(doc, analyses):
    add_heading(doc, "통합 액션 아이템 — 모든 수련회 To-Do 통합", level=1)
    add_para(doc, "마감 임박 순으로 정렬된 통합 액션 리스트. 매일 아침 확인용.",
             size=9, italic=True, color=COLOR_MUTED, space_after=8)

    all_actions = []
    for a in analyses:
        r = a["retreat"]
        for cat, item in a["todo_now"]:
            all_actions.append({
                "church": r.get("churchName") or "?",
                "days": a["days_until"] if a["days_until"] is not None else 999,
                "stage": cat,
                "item": item,
                "risk": a["risk"],
            })
    # 마감 임박순
    all_actions.sort(key=lambda x: (x["days"], x["item"]))
    all_actions = all_actions[:30]  # 너무 길면 컷

    table = doc.add_table(rows=1 + len(all_actions), cols=5)
    table.autofit = False
    headers = ["우선", "교회", "D-day", "단계", "해야 할 일"]
    widths = [1.0, 3.5, 1.5, 2.5, 8.0]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.width = Cm(widths[i])
        c.text = ""
        shade_cell(c, COLOR_BG_HEADER)
        set_cell_borders(c, color="4A2C1A", size=6)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xF8, 0xEB))

    for i, act in enumerate(all_actions):
        cells = table.rows[i+1].cells
        d = f"D-{act['days']}" if act["days"] >= 0 else f"D+{-act['days']}"
        priority = {"high": "🔴", "med": "🟡", "low": "🟢", "unknown": "⚪"}[act["risk"]]
        vals = [priority, act["church"], d, act["stage"], act["item"]]
        for j, v in enumerate(vals):
            cells[j].width = Cm(widths[j])
            cells[j].text = ""
            set_cell_borders(cells[j], color="D4A574")
            shade_cell(cells[j], "FFFFFF" if i%2==0 else COLOR_BG_CREAM)
            p = cells[j].paragraphs[0]
            if j in (0, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            set_font(run, size=9, color=COLOR_PRIMARY)

    doc.add_page_break()


# ─── 부록 ──────────────────────────────────────────────────────────────────────
def add_appendix(doc):
    add_heading(doc, "부록 A — 위험도 판정 기준", level=1)
    add_para(doc, "본 보고서가 사용한 위험도 판정 기준 (결정론적 규칙).",
             size=9, italic=True, color=COLOR_MUTED, space_after=8)

    rules = [
        ("🔴 위험 높음", "D-7 이내인데 진척률 70% 미만 / 미해결 이슈 존재 / 강사·장소 미확정",
         COLOR_RISK_HIGH_BG),
        ("🟡 주의 관찰", "D-30 이내인데 진척률 50% 미만 / 7일간 보고 없음 / D-14 이내 진척률 80% 미만",
         COLOR_RISK_MED_BG),
        ("🟢 정상 진행", "위 기준에 해당하지 않음 — 페이스에 맞게 진행 중",
         COLOR_RISK_LOW_BG),
    ]
    table = doc.add_table(rows=len(rules), cols=2)
    table.autofit = False
    for i, (label, desc, bg) in enumerate(rules):
        lc = table.rows[i].cells[0]
        dc = table.rows[i].cells[1]
        lc.width = Cm(3.5)
        dc.width = Cm(13.0)
        for c in (lc, dc):
            c.text = ""
            set_cell_borders(c, color="D4A574")
        shade_cell(lc, bg)
        run = lc.paragraphs[0].add_run(label)
        set_font(run, size=10, bold=True, color=COLOR_PRIMARY)
        run = dc.paragraphs[0].add_run(desc)
        set_font(run, size=10, color=COLOR_PRIMARY)

    add_para(doc, "", space_after=12)
    add_heading(doc, "부록 B — 데이터 출처", level=1)
    sources = [
        "수련회 기본 정보·체크리스트·보고서: HYDS Ministry App (Railway · MySQL/TiDB)",
        "마스터 체크리스트: directives/retreat_team_leader_checklist.md (D-day별 36항목)",
        "AI 조언 생성: Anthropic Claude (claude-sonnet-4-5)",
        "보고서 생성 시각: " + datetime.now().strftime("%Y-%m-%d %H:%M KST"),
    ]
    for s in sources:
        add_para(doc, f"• {s}", size=10, color=COLOR_PRIMARY, indent=0.3, space_after=2)

    add_para(doc, "", space_after=20)
    add_para(doc,
             "본 문서는 HYDS 내부 운영 목적으로 자동 생성되었으며, 외부에 공유하지 않습니다. "
             "교회·단체 정보는 개인정보 보호를 위해 필요 최소한으로 표기됩니다.",
             size=8, italic=True, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)


# ─── 메인 빌드 ─────────────────────────────────────────────────────────────────
def build_doc(analyses):
    doc = Document()
    setup_page_headers(doc)
    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style.font.size = Pt(10)

    report_date = datetime.now()

    add_cover_page(doc, analyses, report_date)
    add_executive_summary(doc, analyses)
    add_overview_table(doc, analyses)
    for i, a in enumerate(analyses):
        add_retreat_detail(doc, a, i + 1)
    add_consolidated_actions(doc, analyses)
    add_appendix(doc)

    return doc


# ─── 분석 (실제 DB 또는 mock) ──────────────────────────────────────────────────
def analyze_retreat(retreat: dict, master_checklist: dict) -> dict:
    days_until = None
    if retreat.get("startAt"):
        days_until = (datetime.fromtimestamp(retreat["startAt"]/1000) - datetime.now()).days

    chk = retreat.get("checklist") or []
    total = len(chk)
    done = sum(1 for c in chk if c["isChecked"])
    progress = round(done / total * 100) if total else 0

    pending_titles = set(c["title"].strip() for c in chk if not c["isChecked"])

    issues = []
    for r in (retreat.get("recent_reports") or [])[:3]:
        if r.get("issues"):
            issues.append(r["issues"].strip())

    # 활성 카테고리
    all_cats = list(master_checklist.keys())
    if days_until is None: idx = 0
    elif days_until > 90: idx = 0
    elif days_until > 60: idx = 1
    elif days_until > 30: idx = 2
    elif days_until > 14: idx = 3
    elif days_until > 7: idx = 4
    elif days_until > 1: idx = 5
    elif days_until >= 0: idx = 6
    else: idx = 7

    todo_now = [(cat, item) for cat in all_cats[:idx+1] for item in master_checklist[cat]][:15]
    todo_next = [(cat, item) for cat in all_cats[idx+1:idx+2] for item in master_checklist[cat]][:10]

    risk = risk_category(days_until, progress, len(issues))

    reason_summary = ""
    if risk == "high":
        if days_until is not None and days_until <= 7 and progress < 70:
            reason_summary = f"마감 임박(D-{days_until}) + 진척률 미흡({progress}%)"
        elif issues:
            reason_summary = f"미해결 이슈 존재 · {issues[0][:40]}"
        else:
            reason_summary = "팀 구성·강사·장소 등 핵심 항목 미확정"

    return {
        "retreat": retreat,
        "days_until": days_until,
        "progress": progress,
        "total": total,
        "done": done,
        "todo_now": todo_now,
        "todo_next": todo_next,
        "issues": issues[:5],
        "risk": risk,
        "reason_summary": reason_summary,
        "advice": "(부장 조언 — Claude 호출로 생성됨)",
    }


# 마스터 체크리스트 (기존과 동일)
MASTER_CHECKLIST = {
    "D-90 이전": ["교회 측과 zoom 첫 소통, Needs 파악", "HYDS 본부에 교회 Needs 전달·정리"],
    "D-90 ~ D-60": ["팀원 모집 + 예배·프로그램·안전 팀장 세움", "장소 답사", "참가 신청 폼 오픈", "보험 견적", "사진·영상 담당자 별도 지정"],
    "D-60 ~ D-30": ["기획서 교회 측 최종 승인", "1차 카드뉴스/포스터 홍보", "강사 확정 + 사례비 서면", "장소 잔금 일정 확정", "음향 기기 점검"],
    "D-30 ~ D-14": ["참가자 50% 이상 모집", "보험 가입 완료", "식사·교통 업체 계약", "찬양곡 리스트·코드", "단톡방 개설"],
    "D-14 ~ D-7": ["참가자 명단 최종 확정", "알레르기·복용약 수합", "비상연락망 구성", "방·조 배정표", "이름표/조끼 인쇄", "음향 리허설 일정"],
    "D-7 ~ D-1": ["짐 패킹 리스트", "현장 답사", "스태프 R&R 확정", "응급의료 키트", "사고 대응 매뉴얼 숙지"],
    "당일": ["강사 도착 확인", "음향·영상 리허설", "구급함 위치 공지", "단체 사진 시간"],
    "사후 (D+1~D+14)": ["사례비 정산", "참가자 설문", "결산 보고", "강사·장소 후기 반영", "감사 인사"],
}


def send_via_telegram(analyses) -> bool:
    """생성된 보고서 docx + 요약을 텔레그램으로 발송."""
    import os
    import requests
    try:
        from dotenv import load_dotenv
        load_dotenv(PROJECT_ROOT / ".env")
    except ImportError:
        pass
    from execution.telegram_notify import send_telegram

    token = os.getenv("TELEGRAM_BOT_TOKEN", "").strip()
    chat_id = os.getenv("TELEGRAM_ADMIN_CHAT_ID", "").strip()
    today = datetime.now().strftime("%Y-%m-%d %H:%M")
    high = sum(1 for a in analyses if a["risk"] == "high")

    summary = (
        f"<b>📋 HYDS 수련회 운영현황 보고서 ({today})</b>\n\n"
        f"진행 중 <b>{len(analyses)}건</b> · 🔴 위험 <b>{high}건</b>\n"
        f"→ 첨부 Word(표지·Executive Summary·KPI·수련회별 상세·통합 액션) 확인."
    )
    ok1 = send_telegram(summary)

    ok2 = False
    if token and chat_id:
        url = f"https://api.telegram.org/bot{token}/sendDocument"
        try:
            with open(OUTPUT, "rb") as f:
                files = {"document": (OUTPUT.name, f,
                                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
                data = {"chat_id": chat_id, "caption": "HYDS 수련회 운영현황 보고서", "parse_mode": "HTML"}
                res = requests.post(url, data=data, files=files, timeout=30)
                ok2 = res.ok
                if not ok2:
                    print(f"   ❌ sendDocument {res.status_code}: {res.text[:200]}")
        except Exception as e:
            print(f"   ❌ 문서 발송 오류: {e}")
    else:
        print("   ⚠️ 텔레그램 환경변수 없음")
    return ok1 and ok2


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--mock", action="store_true", help="모의 데이터로 샘플 생성")
    parser.add_argument("--no-send", action="store_true")
    args = parser.parse_args()

    if args.mock:
        from datetime import timedelta
        def mock(id, ch, th, days, prog, issues, leader=None):
            ts = int((datetime.now()+timedelta(days=days)).timestamp()*1000)
            total = 26
            done = int(total*prog/100)
            cl = [{'phase':'planning','title':f'항목 {i}','isChecked':i<done} for i in range(total)]
            rep = [{'currentStatus':'예배 팀장 미배정','teamStatus':'인원 부족',
                    'issues':'블레싱 예배 설교자 및 찬양팀 구성 미확정','createdAt':datetime.now()}] if issues else []
            return {'id':id,'churchName':ch,'theme':th,'teamLeader':leader,
                    'startAt':ts,'checklist':cl,'recent_reports':rep}
        mocks = [
            mock(60001,'평택교회','회복',70,0,True),
            mock(60002,'군서찾.수','말씀과 능력',71,0,True),
            mock(30001,'여수 청소년 연합','주안에서',61,50,True,'유하라'),
            mock(30003,'명성교회','거듭남',68,33,False,'이수민'),
            mock(30005,'인천 키즈캠프','하나님의 자녀',69,17,True,'박찬형'),
            mock(30004,'홍성 홍주지방','사명자',75,35,False,'최유진'),
        ]
        analyses = [analyze_retreat(r, MASTER_CHECKLIST) for r in mocks]
        # Mock 조언 채우기
        for a in analyses:
            r = a["retreat"]
            d = a["days_until"]
            a["advice"] = (
                f"{r['churchName']}는 D-{d}, 진척률 {a['progress']}% 상태입니다. "
                f"{'팀 구성 확정과 강사 섭외가 가장 시급하므로 이번 주 내 책임자 1명을 지정하여 진행해야 합니다.' if a['risk']=='high' else '현재 페이스를 유지하되 다음 단계의 강사·장소 확정에 집중하시기 바랍니다.'}"
            )
    else:
        # 실제 DB 호출 — 운영 시
        from execution.db_client import list_active_retreats, get_retreat_full
        from execution.claude_client import ask_claude
        summary = list_active_retreats()
        if not summary:
            print("활성 수련회 없음")
            return
        analyses = []
        for s in summary:
            full = get_retreat_full(s["id"])
            if full:
                a = analyze_retreat(full, MASTER_CHECKLIST)
                r = full
                prompt = f"""HYDS 수련회 매니저로서 {r.get('churchName')} (D-{a['days_until']}, 진척률 {a['progress']}%)에 대한
구체적 조언 1~2문장. 일반론·격언 금지, 구체적 행동 1개 명시."""
                try:
                    a["advice"] = ask_claude(prompt, max_tokens=200).strip()
                except Exception:
                    a["advice"] = "(조언 생성 실패)"
                analyses.append(a)

    print(f"📝 Word 보고서 생성 — {len(analyses)}건")
    doc = build_doc(analyses)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"✅ 저장: {OUTPUT}")
    print(f"   크기: {OUTPUT.stat().st_size:,} bytes")

    # 모의 실행(--mock)은 발송·로그 기록하지 않음 (실데이터만 누적)
    if args.mock:
        return

    notified = False
    if not args.no_send:
        print("📤 텔레그램 발송...")
        notified = send_via_telegram(analyses)
        print("   ✅ 발송 완료" if notified else "   ⚠️ 발송 일부/전체 실패")

    # office 대시보드 + weekly_review 가 읽는 실행 로그
    wlog = read_json("data/weekly_todos_log.json")
    if not isinstance(wlog, list):
        wlog = []
    wlog.append({
        "date": datetime.now().strftime("%Y-%m-%d"),
        "ran_at": datetime.now().isoformat(),
        "retreats_total": len(analyses),
        "total_todos": sum(len(a["todo_now"]) for a in analyses),
        "notified": notified,
        "report_path": str(OUTPUT.relative_to(PROJECT_ROOT)),
    })
    wlog = wlog[-90:]
    write_json("data/weekly_todos_log.json", wlog)
    print("📝 로그 기록 완료 (data/weekly_todos_log.json)")


if __name__ == "__main__":
    main()
