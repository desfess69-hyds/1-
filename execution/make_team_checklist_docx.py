"""
HYDS 팀장 종합 체크리스트 → Word 문서(.docx) 생성.

마크다운 원본을 파싱하지 말고 직접 구조화된 DOCX 생성 (한글 폰트 안전).
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ python-docx 미설치: pip install python-docx")
    sys.exit(1)


PROJECT_ROOT = Path(__file__).parent.parent
OUTPUT = PROJECT_ROOT / ".tmp" / "dossiers" / "HYDS_팀장_종합_체크리스트.docx"
KOREAN_FONT = "Apple SD Gothic Neo"  # Mac 기본. 없으면 Word가 알아서 fallback.


def set_korean_font(run, font_name=KOREAN_FONT):
    """run의 한글 폰트를 명시적으로 설정."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_korean_font(run)
    if level == 0:
        run.font.size = Pt(20)
    elif level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    return h


def add_para(doc, text, bold=False, size=11, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    set_korean_font(run)
    return p


def add_check_list(doc, items):
    """체크리스트 (□ 박스 + 텍스트)."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run("☐  " + item)
        run.font.size = Pt(11)
        set_korean_font(run)


def add_table(doc, header, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, col in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(col)
        run.bold = True
        run.font.size = Pt(11)
        set_korean_font(run)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            set_korean_font(run)
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def build_doc():
    doc = Document()

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr() if hasattr(style.element, 'get_or_add_rPr') else None

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HYDS 수련회 팀장 종합 체크리스트")
    run.font.size = Pt(22)
    run.bold = True
    set_korean_font(run)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("기존 양식 + HYDS 보강 (안전·법·디지털·사후) | 2026-05-29")
    run.font.size = Pt(10)
    run.italic = True
    set_korean_font(run)
    doc.add_paragraph()

    # ─── A. 팀장 5단계 ─────────────────────────
    add_heading(doc, "A. 팀장 5단계 (큰 흐름)", 1)
    add_table(doc, ["단계", "내용", "마감 권장"], [
        ["Step 1", "교회 측과 zoom 첫 소통 — 소통 방식(카톡/문자/줌)·교회 Needs 파악", "D-90"],
        ["Step 2", "HYDS 본부에 교회 Needs 전달·정리", "D-85"],
        ["Step 3", "팀원 모집·업무 배치 (예배 팀장 + 프로그램 팀장 우선)", "D-75"],
        ["Step 4", "답사 (팀장 + 가능한 팀원들 + 기도회 측 교회 챔들)", "D-60"],
        ["Step 5", "세부 기획 회의 — 교회와 지속적 소통", "D-45 ~ D-7"],
    ], col_widths_cm=[2, 11, 3])

    # ─── B. 운영 ─────────────────────────────
    add_heading(doc, "B. 운영 (예배·프로그램 제외 모든 것)", 1)

    add_heading(doc, "B-1. 장소·일정", 2)
    add_check_list(doc, [
        "일정 확정 (1박 2일 / 2박 3일)",
        "숙소 — 스탭 숙소, 캠프 참여자 숙소 별도 확인",
        "온수, 시설(화장실·샤워실 수, WiFi) 확인",
        "식사 — 교회 제공 vs HYDS 도시락",
        "[+] 장소 답사 시 사진·동영상 확보 (팀원 공유용)",
        "[+] 비상 시 대피 동선 사전 확인 (가까운 병원·소방서 거리 메모)",
        "[+] 주차 가능 대수 + 차량 진입 가능 시간",
    ])

    add_heading(doc, "B-2. 캠프 (팀 구성)", 2)
    add_check_list(doc, [
        "팀원 모집 완료",
        "예배 팀장 / 프로그램 팀장 세움",
        "[+] 안전 담당(응급의료) 1명 별도 지정 (간호인 우대)",
        "[+] 사진·영상 담당 1명 별도 지정",
    ])

    add_heading(doc, "B-3. 세부 사항", 2)
    add_check_list(doc, [
        "티셔츠 제작 여부 — 제작 시 디자인 상의 (디자인은 HYDS 제공 가능)",
        "캠프 타임테이블 작성 — 폼은 요청 시 HYDS 제공",
        "캠프 주제 — 교회와 상의 후 결정",
        "[+] 단체 사진 시간 — 시간표에 반영",
    ])

    # ─── C. 예배 ─────────────────────────────
    add_heading(doc, "C. 예배", 1)
    add_para(doc, "핵심: 팀장을 세우고 교회와 상의가 있는 것", italic=True)

    add_heading(doc, "C-1. 설교자", 2)
    add_check_list(doc, [
        "교회와 상의 후 설교자 결정",
        "HYDS 목사님 원할 시 — HYDS 간사님과 소통",
        "[+] 설교자 사례비·교통비·숙소 확정 (서면)",
        "[+] 설교 본문·메시지 방향 사전 공유 받기",
    ])

    add_heading(doc, "C-2. 예배 팀장", 2)
    add_check_list(doc, [
        "인도자 세움 (인도자 = 예배 팀장이 효율적)",
        "예배 큐시트 작성 (요청 시 HYDS 제공)",
        "[+] 찬양곡 리스트 + 코드·악보 사전 공유",
        "[+] 반주자 일정 더블체크",
    ])

    add_heading(doc, "C-3. 음향 기기", 2)
    add_check_list(doc, [
        "예배 시 필요한 음향 기기 파악",
        "교회 기존 장비 외 필요 시 간사와 소통",
        "[+] 마이크·인이어·케이블 여분 (예비 1세트)",
        "[+] 음향 리허설 시간 시간표에 명시",
    ])

    # ─── D. 프로그램 ─────────────────────────
    add_heading(doc, "D. 프로그램", 1)
    add_para(doc, "핵심: 팀장을 세우면 운영이 편함", italic=True)

    add_heading(doc, "D-1. 쉐도우 캠프", 2)
    add_check_list(doc, [
        "쉐도우 캠프 여부 결정",
        "진행 시 — 프로그램 팀장 주관 (프로그램 설명·담당자·장소·진행방법)",
        "[+] 프로그램별 소요시간 실측 (이전 사례 참고)",
    ])

    add_heading(doc, "D-2. 기획서", 2)
    add_check_list(doc, [
        "기획서 작성 (요청 시 HYDS 제공)",
        "[+] 기획서 교회 측 최종 승인 받기 (서면 또는 카톡 확정)",
    ])

    doc.add_page_break()

    # ─── E. HYDS 추가 (큰 강점) ─────────────────
    add_heading(doc, "E. 🆕 HYDS 추가 — 누락되기 쉬운 항목 (66개)", 1)

    add_heading(doc, "E-1. 안전·법·보험", 2)
    add_check_list(doc, [
        "참가자 보험 가입 (50명 이상 필수)",
        "응급의료 키트 준비",
        "가까운 응급실 위치·번호 명단",
        "미성년자 보호자 동의서·연락처 수합",
        "알레르기·복용약·지병 사전 수합",
        "사고 발생 시 대응 매뉴얼 팀장 숙지",
    ])

    add_heading(doc, "E-2. 교통", 2)
    add_check_list(doc, [
        "버스/차량 수배 + 견적서 보관",
        "운전자 명단 + 면허 확인 + 보험 확인",
        "출발 시간·집결 장소 참가자 공지",
        "귀가 시 픽업 동선 미성년자 보호자 안내",
    ])

    add_heading(doc, "E-3. 참가자 관리", 2)
    add_check_list(doc, [
        "참가 신청서 양식 (이름·연락처·소속·알레르기·비상연락)",
        "마감일 기준 참가자 명단 확정",
        "단톡방 개설 + 공지 매뉴얼",
        "방·조 배정표 사전 공유",
        "이름표/조끼 미리 인쇄",
    ])

    add_heading(doc, "E-4. 예산·정산", 2)
    add_check_list(doc, [
        "항목별 예산표 교회와 합의 (서면)",
        "사례비·강사비·장소비 정산 일정 명시",
        "영수증 보관 담당자 지정",
        "계좌·법인카드 사용 계획",
        "사후 결산 보고서 D+7 이내 제출",
    ])

    add_heading(doc, "E-5. 홍보·모집", 2)
    add_check_list(doc, [
        "신청 폼 오픈 (D-60)",
        "카드뉴스/포스터 D-45 1차, D-15 마감 임박 푸쉬",
        "단톡방 매일 D-day 공지",
        "모집 인원 미달 시 비상 플랜 (대상 확대 or 일정 축소)",
    ])

    add_heading(doc, "E-6. 기록·동의", 2)
    add_check_list(doc, [
        "사진·영상 촬영 동의서 (미성년자는 보호자 동의)",
        "개인정보 수집·이용 동의서",
        "수련회 사진·영상 활용 동의 (SNS·홍보용)",
        "드라이브 폴더 사진/영상 업로드 약속",
    ])

    add_heading(doc, "E-7. 디지털·체크인", 2)
    add_check_list(doc, [
        "출석 체크 방식 결정 (수기 / 단톡 / QR)",
        "시간표 자동 알림 (단톡 봇 or 카톡 알림장)",
        "분실물·습득물 단톡방 채널",
    ])

    add_heading(doc, "E-8. 사후 (D+1 ~ D+14)", 2)
    add_check_list(doc, [
        "참가자 설문 발송 (D+3 이내)",
        "결산 보고 교회 측에 (D+7 이내)",
        "강사·장소·진행 후기 HYDS 데이터에 반영",
        "다음 기획에 적용할 교훈 메모",
        "참가자·교회 감사 인사",
    ])

    doc.add_page_break()

    # ─── F. D-day 요약 ─────────────────────
    add_heading(doc, "F. 점검 시점별 핵심 (요약)", 1)
    add_table(doc, ["D-day", "무조건 확인"], [
        ["D-90", "교회 첫 미팅 / Needs 파악"],
        ["D-75", "팀원 + 팀장(예배·프로그램·안전) 확정"],
        ["D-60", "답사 / 신청 폼 오픈 / 보험 견적"],
        ["D-45", "기획서 교회 승인 / 1차 홍보"],
        ["D-30", "강사·장소 잔금 일정 / 보험 가입 / 50% 모집"],
        ["D-14", "식·교통 계약 확정 / 음향 리허설 일정"],
        ["D-7", "명단 확정 / 알레르기·약 수합 / 비상연락망"],
        ["D-1", "짐 패킹 / 현장 답사 / 스태프 R&R"],
        ["당일", "강사 도착 / 리허설 / 응급키트 위치 공지"],
        ["D+1 ~ D+14", "설문·결산·강사 후기 / 학습 내용 반영"],
    ], col_widths_cm=[3, 13])

    # 푸터
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run("이 문서는 HYDS의 살아있는 자산입니다. 매 수련회 후 학습 내용을 추가하세요.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    set_korean_font(run)

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUTPUT)
    print(f"✅ Word 파일 생성: {OUTPUT}")
    print(f"   파일 크기: {OUTPUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
